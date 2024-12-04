from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

class WordGenerator:
    def __init__(self):
        self.document = Document()
        self._setup_document()

    def _setup_document(self):
        """Configurar el estilo del documento"""
        # Configurar márgenes
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

    def add_title(self, title):
        """Agregar título principal"""
        heading = self.document.add_heading(title, level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Agregar espacio después del título
        self.document.add_paragraph()

    def add_process_section(self, title, content):
        """Agregar una sección del proceso"""
        # Agregar título de sección
        heading = self.document.add_heading(title, level=1)
        heading.style.font.size = Pt(14)
        
        # Agregar contenido
        if isinstance(content, list):
            for item in content:
                p = self.document.add_paragraph()
                p.add_run(item)
        else:
            p = self.document.add_paragraph()
            p.add_run(content)

    def add_subprocess_steps(self, steps):
        """Agregar pasos del subproceso"""
        for i, step in enumerate(steps, 1):
            p = self.document.add_paragraph()
            p.add_run(f"{i}. {step}")

    def save(self, filepath):
        """Guardar el documento"""
        self.document.save(filepath)

def generate_process_document(process_name, process_description, subprocesses):
    """
    Generar documento Word con la descripción del proceso
    
    Args:
        process_name (str): Nombre del proceso
        process_description (str): Descripción del proceso
        subprocesses (list): Lista de subprocesos con sus descripciones
    """
    generator = WordGenerator()
    
    # Agregar título del proceso
    generator.add_title(f"Proceso: {process_name}")
    
    # Agregar descripción del proceso
    generator.add_process_section("1. ALCANCE DEL PROCESO", process_description)
    
    # Agregar desarrollo del proceso
    generator.add_process_section("2. DESARROLLO DEL PROCESO", "")
    
    # Agregar subprocesos
    for i, subprocess in enumerate(subprocesses, 1):
        generator.add_process_section(
            f"2.{i}. {subprocess['name']}", 
            subprocess['description']
        )
    
    # Agregar criterios de control si existen
    if any('control' in sub for sub in subprocesses):
        generator.add_process_section(
            "3. CRITERIOS DE CONTROL",
            [sub['control'] for sub in subprocesses if 'control' in sub]
        )
    
    # Crear directorio para documentos si no existe
    docs_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs')
    os.makedirs(docs_dir, exist_ok=True)
    
    # Guardar documento
    filepath = os.path.join(docs_dir, f"{process_name.replace(' ', '_')}.docx")
    generator.save(filepath)
    
    return filepath
